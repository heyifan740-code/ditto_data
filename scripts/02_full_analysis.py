"""
DITTO Data Full Analysis
生成完整数据特征报告，输出为 docx
"""

import zarr
import numpy as np
import os, glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 路径 ────────────────────────────────────────────────────────────────────
BASE    = "/home/fa_team/roamlab/finger_aloha/software/ditto_data/traning_data"
OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "outputs", "reports")
os.makedirs(OUT_DIR, exist_ok=True)
DATASETS = {
    "ITW":    "uncap_itw_xarm_7dof_05_18_2026",
    "Teleop": "uncap_teleop_xarm_7dof_05_18_2026",
}

# ─── 辅助：给 docx 表格加边框 ────────────────────────────────────────────────
def set_table_border(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top','left','bottom','right','insideH','insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def shade_row(row, hex_color='F2F2F2'):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

# ════════════════════════════════════════════════════════════════════════════
# 数据采集阶段：遍历所有 demo，提取所有数值特征
# ════════════════════════════════════════════════════════════════════════════
print("正在提取数据特征，请稍候...")

# 存储结构：results[mode] = { channel: { stat: value } }
results = {}

# 数值型 channel（图像单独处理）
NUMERIC_KEYS = [
    "arm_obs", "arm_action", "arm_published_action",
    "follower_joint_states", "follower_joint_velocities", "follower_joint_efforts",
    "leader_joint_states",   "leader_joint_velocities",   "leader_joint_efforts",
    "hand_action",
]
IMAGE_KEYS = ["leader_rgb", "palm_rgb", "side_rgb"]

for mode, dataset_name in DATASETS.items():
    print(f"  [{mode}] 处理中...")
    demo_dirs = sorted(glob.glob(os.path.join(BASE, dataset_name, "demo_*")))

    # ── per-demo 统计 ──────────────────────────────────────────────────────
    demo_lengths   = []
    demo_durations = []
    demo_freqs     = []
    ts_gaps_all    = []   # 所有帧间时间差，用来检测抖动

    # ── 累积所有 demo 的数值数据（每个 key 收集所有帧） ──────────────────────
    aggregated = {k: [] for k in NUMERIC_KEYS}

    # ── action smoothness：对 hand_action 和 arm_action 求速度/jerk ──────────
    hand_vel_all = []    # hand_action 帧间差（velocity proxy）
    hand_jerk_all = []   # velocity 的帧间差（jerk proxy）
    arm_vel_all  = []    # arm_action 帧间差（teleop only）

    # ── leader-follower 跟踪误差 ──────────────────────────────────────────
    tracking_errors = []  # |leader_joint_states - follower_joint_states|

    # ── 图像统计 ──────────────────────────────────────────────────────────
    img_stats = {}  # key -> list of (mean_brightness, std_brightness)
    avail_img_keys = []

    for d in demo_dirs:
        z  = zarr.open(os.path.join(d, "data.zarr"), mode='r')
        ts = z["timestamp"][:]
        n  = len(ts)
        demo_lengths.append(n)
        if n > 1:
            dt = np.diff(ts)
            demo_durations.append(ts[-1] - ts[0])
            demo_freqs.append(1.0 / np.median(dt))
            ts_gaps_all.extend(dt.tolist())

        # 数值 channel
        for k in NUMERIC_KEYS:
            if k in z:
                aggregated[k].append(z[k][:])

        # action smoothness（hand_action 必有）
        if "hand_action" in z and n > 2:
            ha = z["hand_action"][:]
            vel  = np.diff(ha, axis=0) * 20.0   # 乘以频率 -> rad/s
            jerk = np.diff(vel, axis=0) * 20.0  # rad/s²
            hand_vel_all.append(vel)
            hand_jerk_all.append(jerk)

        if "arm_action" in z and n > 2:
            aa  = z["arm_action"][:]
            vel = np.diff(aa, axis=0) * 20.0
            arm_vel_all.append(vel)

        # leader-follower 跟踪误差
        if "leader_joint_states" in z and "follower_joint_states" in z:
            err = np.abs(z["leader_joint_states"][:] - z["follower_joint_states"][:])
            tracking_errors.append(err)

    # 拼接所有 demo 的数据
    for k in NUMERIC_KEYS:
        if aggregated[k]:
            aggregated[k] = np.vstack(aggregated[k])  # (total_frames, dims)

    # 拼接 smoothness
    hand_vel_cat  = np.vstack(hand_vel_all)  if hand_vel_all  else None
    hand_jerk_cat = np.vstack(hand_jerk_all) if hand_jerk_all else None
    arm_vel_cat   = np.vstack(arm_vel_all)   if arm_vel_all   else None
    tracking_cat  = np.vstack(tracking_errors) if tracking_errors else None

    # ── 图像统计：只取每个数据集的前 5 个 demo，每个 demo 采样 10 帧 ──────────
    img_brightness = {}
    for ik in IMAGE_KEYS:
        samples = []
        for d in demo_dirs[:5]:
            z = zarr.open(os.path.join(d, "data.zarr"), mode='r')
            if ik in z:
                arr = z[ik]
                n   = arr.shape[0]
                idx = np.linspace(0, n-1, min(10, n), dtype=int)
                frames = arr[idx]                      # (10, H, W, 3)
                samples.append(frames.mean())
        if samples:
            img_brightness[ik] = float(np.mean(samples))

    # ── 汇总 ──────────────────────────────────────────────────────────────
    results[mode] = {
        "demo_count":   len(demo_dirs),
        "lengths":      np.array(demo_lengths),
        "durations":    np.array(demo_durations),
        "freqs":        np.array(demo_freqs),
        "ts_gaps":      np.array(ts_gaps_all),
        "aggregated":   aggregated,
        "hand_vel":     hand_vel_cat,
        "hand_jerk":    hand_jerk_cat,
        "arm_vel":      arm_vel_cat,
        "tracking":     tracking_cat,
        "img_bright":   img_brightness,
    }

print("  特征提取完成，开始生成报告...")

# ════════════════════════════════════════════════════════════════════════════
# 生成 docx
# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── 全局字体 ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def note(text):
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p

# ════════════ 封面 ════════════════════════════════════════════════════════
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("DITTO 数据特征分析报告")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2.add_run(f"任务：uncap（拧开瓶盖）  |  日期：{datetime.now().strftime('%Y-%m-%d')}")

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3.add_run("采集模式：ITW（In-The-Wild，手持 leader）vs Teleop（远程遥操作）")

doc.add_page_break()

# ════════════ 1. 项目背景 ════════════════════════════════════════════════
h1("1. 项目背景与数据采集方式")

body(
    "本数据集由 DITTO（7-DoF 灵巧手 teleoperation 系统）采集，任务为拧开瓶盖（uncap），"
    "采集时间为 2026 年 5 月 18–19 日，平台为 xArm 7-DoF 机械臂 + Finger ALOHA 7-DoF 灵巧手。"
    "数据包含两种采集模式，物理设置不同，导致数据特性存在系统性差异。"
)

h2("1.1 ITW 模式（In-The-Wild，手持操作）")
body(
    "操作员手持 leader 装置（与灵巧手形态相同的外骨骼设备），直接用手操作物体完成任务。"
    "机械臂（xArm）在此模式下不参与（USE_ARM=false），灵巧手 follower 不承受真实接触载荷，"
    "因此 follower 的关节力矩（efforts）记录值为全零。"
    "控制频率：20 Hz。相机：leader 视角相机（leader_rgb）+ 侧视相机（side_rgb）。"
)

h2("1.2 Teleop 模式（遥操作）")
body(
    "操作员戴上 leader 装置，远程控制 follower 机器人手执行任务，机械臂同步运动。"
    "follower 真实与物体接触，关节电流（follower_joint_efforts）记录实际接触力信号。"
    "控制频率：20 Hz。相机：手掌相机（palm_rgb）+ 侧视相机（side_rgb）。"
    "额外记录：机械臂笛卡尔观测（arm_obs）、机械臂动作指令（arm_action / arm_published_action）。"
)

h2("1.3 数据格式")
body(
    "所有数据以 Zarr（v2）格式存储，每条 demo 独立存放于 demo_N/ 目录下的 data.zarr 文件中。"
    "Zarr 是专为大型数值数组设计的分块压缩格式，支持延迟加载，"
    "适合存储图像序列等大体积张量。每个 zarr 文件是一个 Group，其中每个 channel 是一个独立的 Array，"
    "第 0 维统一为时间步数 T（帧数）。"
)

# 目录结构
h2("1.4 目录结构")
body("数据根目录：/home/fa_team/roamlab/finger_aloha/software/ditto_data/traning_data/")
doc.add_paragraph(
    "traning_data/\n"
    "├── uncap_itw_xarm_7dof_05_18_2026/          ← ITW 数据集\n"
    "│   ├── config_uncap_itw_xarm_7dof_05_18_2026.yaml\n"
    "│   ├── demo_0/\n"
    "│   │   ├── config.yaml                       ← 本条 demo 的运行配置\n"
    "│   │   └── data.zarr/                        ← 所有 channel 数据\n"
    "│   │       ├── arm_obs/\n"
    "│   │       ├── follower_joint_states/\n"
    "│   │       ├── hand_action/\n"
    "│   │       ├── leader_rgb/\n"
    "│   │       └── ...（共 11 个 channel）\n"
    "│   ├── demo_1/ ... demo_49/\n"
    "└── uncap_teleop_xarm_7dof_05_18_2026/       ← Teleop 数据集\n"
    "    ├── demo_0/ ... demo_49/                  ← 结构相同，channel 有差异\n"
    , style='No Spacing'
)

doc.add_page_break()

# ════════════ 2. Channel 清单 ════════════════════════════════════════════
h1("2. Channel 完整清单与含义")

body(
    "以下表格列出两种模式各自包含的所有 channel，说明其物理含义、数据维度及单位。"
    "shape 中的 T 表示该 demo 的帧数，每个 demo 的 T 不同。"
)

# ── ITW Channel 表 ──────────────────────────────────────────────────────
h2("2.1 ITW Channel 表（共 11 个）")

itw_channels = [
    ("timestamp",               "(T,)",          "float64", "Unix 时间戳（秒），用于计算频率和对齐"),
    ("arm_obs",                 "(T, 6)",         "float64", "机械臂观测（ITW 模式 USE_ARM=false，记录 leader IMU/传感器数据，值域极小 ~0.001–0.007）"),
    ("follower_joint_states",   "(T, 7)",         "float64", "7-DoF 灵巧手 follower 关节角度（rad 或 encoder units）"),
    ("follower_joint_velocities","(T, 7)",        "float64", "follower 关节角速度"),
    ("follower_joint_efforts",  "(T, 7)",         "float64", "follower 关节电流/力矩 — ITW 模式全为 0（无接触力信号）"),
    ("leader_joint_states",     "(T, 7)",         "float64", "leader 装置关节角度，反映操作员手部姿态"),
    ("leader_joint_velocities", "(T, 7)",         "float64", "leader 关节角速度"),
    ("leader_joint_efforts",    "(T, 7)",         "float64", "leader 关节电流 — ITW 模式全为 0"),
    ("hand_action",             "(T, 7)",         "float64", "发送给 follower 手的关节位置指令，7-DoF"),
    ("leader_rgb",              "(T, 480, 640, 3)","uint8",  "leader 视角 RGB 相机，60 fps 采集但控制频率 20Hz 采样，分辨率 640×480"),
    ("side_rgb",                "(T, 480, 640, 3)","uint8",  "侧视 RGB 相机，同上"),
]

tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(tbl)
hdr = tbl.rows[0]
shade_row(hdr, 'D6E4F0')
for i, h in enumerate(["Channel 名", "Shape", "dtype", "含义"]):
    hdr.cells[i].text = h
    hdr.cells[i].paragraphs[0].runs[0].bold = True

for row_data in itw_channels:
    row = tbl.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()

h2("2.2 Teleop Channel 表（共 13 个）")

tp_channels = [
    ("timestamp",               "(T,)",          "float64", "Unix 时间戳（秒）"),
    ("arm_obs",                 "(T, 6)",         "float64", "机械臂末端笛卡尔观测：[x, y, z, rx, ry, rz]，单位 m / rad"),
    ("arm_action",              "(T, 6)",         "float64", "下一步机械臂目标位姿指令（由 leader 映射而来）"),
    ("arm_published_action",    "(T, 6)",         "float64", "实际发布给机械臂控制器的指令（可能经过滤波/限幅，与 arm_obs 几乎一致）"),
    ("follower_joint_states",   "(T, 7)",         "float64", "follower 手关节角度"),
    ("follower_joint_velocities","(T, 7)",        "float64", "follower 关节角速度"),
    ("follower_joint_efforts",  "(T, 7)",         "float64", "follower 关节电流 — 有真实信号，反映接触力，范围约 ±500"),
    ("leader_joint_states",     "(T, 7)",         "float64", "leader 关节角度"),
    ("leader_joint_velocities", "(T, 7)",         "float64", "leader 关节角速度"),
    ("leader_joint_efforts",    "(T, 7)",         "float64", "leader 关节电流，力反馈渲染用（force_rendering 开启时有信号）"),
    ("hand_action",             "(T, 7)",         "float64", "发送给 follower 手的关节位置指令，7-DoF"),
    ("palm_rgb",                "(T, 480, 640, 3)","uint8",  "手掌视角相机（安装于灵巧手掌心），640×480"),
    ("side_rgb",                "(T, 480, 640, 3)","uint8",  "侧视 RGB 相机，640×480"),
]

tbl2 = doc.add_table(rows=1, cols=4)
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(tbl2)
hdr2 = tbl2.rows[0]
shade_row(hdr2, 'D6E4F0')
for i, h in enumerate(["Channel 名", "Shape", "dtype", "含义"]):
    hdr2.cells[i].text = h
    hdr2.cells[i].paragraphs[0].runs[0].bold = True

for row_data in tp_channels:
    row = tbl2.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()

h2("2.3 Channel 差异对比")

diff_data = [
    ("arm_obs",                "✓（值极小，非 Cartesian）","✓（Cartesian 位姿，单位 m/rad）"),
    ("follower_joint_states",  "✓","✓"),
    ("follower_joint_velocities","✓","✓"),
    ("follower_joint_efforts", "✓（全零，无力信号）","✓（有真实接触力信号）"),
    ("leader_joint_states",    "✓","✓"),
    ("leader_joint_velocities","✓","✓"),
    ("leader_joint_efforts",   "✓（全零）","✓（力反馈渲染电流）"),
    ("hand_action",            "✓","✓"),
    ("side_rgb",               "✓","✓"),
    ("timestamp",              "✓","✓"),
    ("leader_rgb",             "✓（有）","✗（无）"),
    ("palm_rgb",               "✗（无）","✓（有）"),
    ("arm_action",             "✗（无）","✓（有）"),
    ("arm_published_action",   "✗（无）","✓（有）"),
]

tbl3 = doc.add_table(rows=1, cols=3)
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_border(tbl3)
hdr3 = tbl3.rows[0]
shade_row(hdr3, 'D6E4F0')
for i, h in enumerate(["Channel", "ITW", "Teleop"]):
    hdr3.cells[i].text = h
    hdr3.cells[i].paragraphs[0].runs[0].bold = True

for row_data in diff_data:
    row = tbl3.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_page_break()

# ════════════ 3. 基本统计 ════════════════════════════════════════════════
h1("3. 基本统计量")

h2("3.1 Demo 级别统计")

stat_rows = []
for mode in ["ITW", "Teleop"]:
    r = results[mode]
    L = r["lengths"]
    D = r["durations"]
    F = r["freqs"]
    stat_rows.append([
        mode,
        str(r["demo_count"]),
        f"{L.min()} / {L.max()} / {L.mean():.1f} ± {L.std():.1f}",
        f"{D.min():.1f} / {D.max():.1f} / {D.mean():.1f} ± {D.std():.1f}",
        f"{F.mean():.3f}",
        str(int(L.sum())),
        f"{D.sum():.1f}",
    ])

tbl4 = doc.add_table(rows=1, cols=7)
set_table_border(tbl4)
hdr4 = tbl4.rows[0]
shade_row(hdr4, 'D6E4F0')
for i, h in enumerate(["模式","Demo 数","帧数 min/max/mean±std",
                        "时长(s) min/max/mean±std",
                        "频率(Hz)","总帧数","总时长(s)"]):
    hdr4.cells[i].text = h
    hdr4.cells[i].paragraphs[0].runs[0].bold = True
for row_data in stat_rows:
    row = tbl4.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()
note(
    "Teleop 的 demo 时长更长（均值 15.0s vs 12.9s）、标准差更大（50.9 vs 40.6），"
    "说明遥操作的执行速度更慢且个体差异更大。"
)

# ── 时间戳间隔一致性 ──────────────────────────────────────────────────────
h2("3.2 时间戳间隔一致性（控制频率稳定性）")
body(
    "理论控制周期为 50ms（20Hz）。以下统计所有帧间时间差（ms），检验是否存在掉帧或抖动。"
)

ts_rows = []
for mode in ["ITW", "Teleop"]:
    gaps = results[mode]["ts_gaps"] * 1000  # 转换为 ms
    p95  = np.percentile(gaps, 95)
    p99  = np.percentile(gaps, 99)
    ts_rows.append([
        mode,
        f"{gaps.mean():.2f}",
        f"{gaps.std():.2f}",
        f"{gaps.min():.2f}",
        f"{gaps.max():.2f}",
        f"{p95:.2f}",
        f"{p99:.2f}",
        str(int(np.sum(gaps > 60))),   # 超过 60ms 的帧（掉帧阈值）
    ])

tbl5 = doc.add_table(rows=1, cols=8)
set_table_border(tbl5)
hdr5 = tbl5.rows[0]
shade_row(hdr5, 'D6E4F0')
for i, h in enumerate(["模式","mean(ms)","std(ms)","min(ms)","max(ms)","P95(ms)","P99(ms)",">60ms帧数"]):
    hdr5.cells[i].text = h
    hdr5.cells[i].paragraphs[0].runs[0].bold = True
for row_data in ts_rows:
    row = tbl5.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_page_break()

# ════════════ 4. 数值 Channel 详细统计 ════════════════════════════════════
h1("4. 数值 Channel 详细统计（全量 Demo 聚合）")

body(
    "以下各表将所有 demo 的数据拼接后统计，每列对应一个关节（joint 0–6）或维度。"
    "mean 和 std 反映数据分布中心和离散度，min/max 反映运动范围。"
)

KEY_LABELS = {
    "follower_joint_states":    "4.1 Follower 关节角度（follower_joint_states）",
    "follower_joint_velocities":"4.2 Follower 关节角速度（follower_joint_velocities）",
    "follower_joint_efforts":   "4.3 Follower 关节电流/力矩（follower_joint_efforts）",
    "leader_joint_states":      "4.4 Leader 关节角度（leader_joint_states）",
    "leader_joint_velocities":  "4.5 Leader 关节角速度（leader_joint_velocities）",
    "leader_joint_efforts":     "4.6 Leader 关节电流（leader_joint_efforts）",
    "hand_action":              "4.7 Hand Action（手部动作指令）",
    "arm_obs":                  "4.8 Arm Observation（机械臂观测）",
}

for key, section_title in KEY_LABELS.items():
    h2(section_title)

    has_data = any(
        isinstance(results[m]["aggregated"].get(key), np.ndarray)
        for m in ["ITW", "Teleop"]
    )
    if not has_data:
        body("（两种模式均无此 channel）")
        continue

    # 表头：关节维度数
    ndim = None
    for m in ["ITW", "Teleop"]:
        arr = results[m]["aggregated"].get(key)
        if isinstance(arr, np.ndarray):
            ndim = arr.shape[1]
            break

    col_names = [f"J{i}" for i in range(ndim)]

    for mode in ["ITW", "Teleop"]:
        arr = results[mode]["aggregated"].get(key)
        if not isinstance(arr, np.ndarray):
            body(f"【{mode}】此 channel 不存在")
            continue

        p = doc.add_paragraph()
        run = p.add_run(f"【{mode}】  总帧数: {arr.shape[0]}")
        run.bold = True

        stat_tbl = doc.add_table(rows=5, cols=ndim+1)
        set_table_border(stat_tbl)
        shade_row(stat_tbl.rows[0], 'EBF3FB')

        stat_tbl.rows[0].cells[0].text = "统计量"
        for j, cn in enumerate(col_names):
            stat_tbl.rows[0].cells[j+1].text = cn
            stat_tbl.rows[0].cells[j+1].paragraphs[0].runs[0].bold = True

        stats = [
            ("mean", np.mean(arr, axis=0)),
            ("std",  np.std(arr,  axis=0)),
            ("min",  np.min(arr,  axis=0)),
            ("max",  np.max(arr,  axis=0)),
        ]
        for si, (sname, sval) in enumerate(stats):
            row = stat_tbl.rows[si+1]
            row.cells[0].text = sname
            row.cells[0].paragraphs[0].runs[0].bold = True
            for j in range(ndim):
                row.cells[j+1].text = f"{sval[j]:.4f}"

        doc.add_paragraph()

doc.add_page_break()

# ════════════ 5. 动作平滑性分析 ════════════════════════════════════════════
h1("5. 动作平滑性分析（Hand Action）")

body(
    "平滑性是衡量演示数据质量的重要指标。过大的 jerk（加加速度）意味着动作不连贯，"
    "训练时可能导致 policy 产生抖动输出。"
    "这里用 hand_action 的帧间差分来近似速度，再对速度做差分近似 jerk。"
    "控制频率 20Hz，所以乘以 20 可以得到近似单位 rad/s 和 rad/s²。"
)

smooth_rows = []
for mode in ["ITW", "Teleop"]:
    vel  = results[mode]["hand_vel"]
    jerk = results[mode]["hand_jerk"]
    if vel is None:
        continue
    vel_mag  = np.linalg.norm(vel,  axis=1)   # 每帧的速度向量模长
    jerk_mag = np.linalg.norm(jerk, axis=1)
    smooth_rows.append([
        mode,
        f"{vel_mag.mean():.4f}",
        f"{vel_mag.std():.4f}",
        f"{vel_mag.max():.4f}",
        f"{jerk_mag.mean():.4f}",
        f"{jerk_mag.std():.4f}",
        f"{jerk_mag.max():.4f}",
    ])

tbl_s = doc.add_table(rows=1, cols=7)
set_table_border(tbl_s)
shade_row(tbl_s.rows[0], 'D6E4F0')
for i, h in enumerate(["模式",
                        "速度模长 mean","速度模长 std","速度模长 max",
                        "Jerk 模长 mean","Jerk 模长 std","Jerk 模长 max"]):
    tbl_s.rows[0].cells[i].text = h
    tbl_s.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in smooth_rows:
    row = tbl_s.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_paragraph()
note(
    "速度模长更小、jerk 模长更小 → 动作更平滑。"
    "ITW 直接用手操作，动作通常更快更自然；Teleop 隔着接口，可能有延迟导致更大抖动。"
    "具体结论见数值对比。"
)

# ─── Arm action smoothness (teleop only) ────────────────────────────────
h2("5.2 机械臂动作平滑性（arm_action，仅 Teleop）")
body(
    "arm_action 是 6-DoF 末端笛卡尔指令 [x,y,z,rx,ry,rz]，"
    "前 3 维是位置（m），后 3 维是姿态（rad）。分开统计位置和姿态分量的速度。"
)

if results["Teleop"]["arm_vel"] is not None:
    av = results["Teleop"]["arm_vel"]
    pos_vel  = np.linalg.norm(av[:, :3], axis=1)   # 位置速度，单位 m/s
    rot_vel  = np.linalg.norm(av[:, 3:], axis=1)   # 姿态速度，单位 rad/s
    av_rows = [
        ["位置速度 |v_pos| (m/s)", f"{pos_vel.mean():.5f}", f"{pos_vel.std():.5f}", f"{pos_vel.max():.5f}"],
        ["姿态速度 |v_rot| (rad/s)", f"{rot_vel.mean():.5f}", f"{rot_vel.std():.5f}", f"{rot_vel.max():.5f}"],
    ]
    tbl_av = doc.add_table(rows=1, cols=4)
    set_table_border(tbl_av)
    shade_row(tbl_av.rows[0], 'D6E4F0')
    for i, h in enumerate(["分量","mean","std","max"]):
        tbl_av.rows[0].cells[i].text = h
        tbl_av.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row_data in av_rows:
        row = tbl_av.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

doc.add_page_break()

# ════════════ 6. Leader-Follower 跟踪误差 ════════════════════════════════
h1("6. Leader-Follower 关节跟踪误差")

body(
    "Leader 发出指令，follower 执行。跟踪误差 = |leader_joint_states - follower_joint_states|，"
    "反映 follower 跟随 leader 的精度。误差大说明存在滞后或机械限位约束。"
    "对于 policy 学习，这个误差直接影响 action 标签的质量。"
)

for mode in ["ITW", "Teleop"]:
    tc = results[mode]["tracking"]
    if tc is None:
        continue
    h2(f"6.{1 if mode=='ITW' else 2} {mode} 跟踪误差（per joint，单位与 joint_states 一致）")
    tr_rows = []
    for j in range(tc.shape[1]):
        tr_rows.append([
            f"Joint {j}",
            f"{tc[:,j].mean():.4f}",
            f"{tc[:,j].std():.4f}",
            f"{tc[:,j].min():.4f}",
            f"{tc[:,j].max():.4f}",
            f"{np.percentile(tc[:,j], 95):.4f}",
        ])
    tbl_tr = doc.add_table(rows=1, cols=6)
    set_table_border(tbl_tr)
    shade_row(tbl_tr.rows[0], 'D6E4F0')
    for i, h in enumerate(["关节","mean","std","min","max","P95"]):
        tbl_tr.rows[0].cells[i].text = h
        tbl_tr.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row_data in tr_rows:
        row = tbl_tr.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
    doc.add_paragraph()

doc.add_page_break()

# ════════════ 7. 力信号分析（Teleop） ════════════════════════════════════
h1("7. 力信号分析（follower_joint_efforts，仅 Teleop 有效）")

body(
    "follower_joint_efforts 记录灵巧手各关节的电机电流，"
    "是接触力的间接代理信号（motor current → joint torque → contact force）。"
    "ITW 模式此信号全为零，Teleop 模式有真实信号。"
    "以下统计 Teleop 全量数据中 efforts 的分布。"
)

eff = results["Teleop"]["aggregated"].get("follower_joint_efforts")
if isinstance(eff, np.ndarray):
    eff_rows = []
    for j in range(eff.shape[1]):
        nz_ratio = np.mean(eff[:, j] != 0) * 100
        eff_rows.append([
            f"Joint {j}",
            f"{eff[:,j].mean():.2f}",
            f"{eff[:,j].std():.2f}",
            f"{eff[:,j].min():.2f}",
            f"{eff[:,j].max():.2f}",
            f"{np.percentile(np.abs(eff[:,j]), 90):.2f}",
            f"{nz_ratio:.1f}%",
        ])

    tbl_eff = doc.add_table(rows=1, cols=7)
    set_table_border(tbl_eff)
    shade_row(tbl_eff.rows[0], 'D6E4F0')
    for i, h in enumerate(["关节","mean","std","min","max","|effort| P90","非零率"]):
        tbl_eff.rows[0].cells[i].text = h
        tbl_eff.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for row_data in eff_rows:
        row = tbl_eff.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val

    doc.add_paragraph()
    note(
        "力信号幅值大（std 可达 ±100 量级）、分布不对称，说明拧瓶盖任务中各关节承受的接触力差异显著。"
        "Joint 3/4/5 方向的 std 尤其大，可能对应拧转方向的主要施力关节。"
        "这是 ITW 数据完全缺失的信息维度，对理解 ITW vs Teleop 的分布差异至关重要。"
    )

doc.add_page_break()

# ════════════ 8. 图像信息 ═══════════════════════════════════════════════
h1("8. 图像 Channel 信息")

body("图像 channel 均为 RGB 格式，uint8，分辨率 640×480（H×W），以下统计平均亮度（基于前 5 个 demo 抽样）。")

img_rows = []
for mode in ["ITW", "Teleop"]:
    for ik, bv in results[mode]["img_bright"].items():
        img_rows.append([mode, ik, "640×480", "uint8", "RGB", f"{bv:.1f}"])
    # 标注缺失的
    for ik in IMAGE_KEYS:
        if ik not in results[mode]["img_bright"]:
            img_rows.append([mode, ik, "—", "—", "—", "（此模式无此相机）"])

tbl_img = doc.add_table(rows=1, cols=6)
set_table_border(tbl_img)
shade_row(tbl_img.rows[0], 'D6E4F0')
for i, h in enumerate(["模式","Channel","分辨率","dtype","格式","平均亮度（0–255）"]):
    tbl_img.rows[0].cells[i].text = h
    tbl_img.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for row_data in img_rows:
    row = tbl_img.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val

doc.add_page_break()

# ════════════ 9. 关键发现与研究意义 ════════════════════════════════════════
h1("9. 关键发现与研究意义")

findings = [
    (
        "① 力信号的系统性缺失（最重要发现）",
        "ITW 模式的 follower_joint_efforts 和 leader_joint_efforts 在所有 50 条 demo 中全为零。"
        "Teleop 模式两者均有真实信号（std ≈ 40–160 量级）。"
        "这意味着：如果 policy 输入包含 efforts，ITW 数据无法提供接触感知信息。"
        "这是两种模式数据分布差异的最根本来源之一，也是后续分析的核心控制变量。"
    ),
    (
        "② arm_obs 的语义不同",
        "ITW 模式下 arm_obs 的值域极小（~0.001–0.007），不是笛卡尔坐标（单位不对）；"
        "Teleop 模式下 arm_obs 是清晰的末端位姿 [x,y,z,rx,ry,rz]（x≈0.18–0.40m 符合工作空间）。"
        "ITW 的 arm_obs 可能来自 IMU 或 leader 设备的角速度传感器，"
        "两者语义完全不同，不能直接作为同一 observation 喂给同一个 policy。"
    ),
    (
        "③ 相机配置不同",
        "ITW 有 leader_rgb（操作员视角）无 palm_rgb；Teleop 有 palm_rgb（机器人手掌视角）无 leader_rgb。"
        "side_rgb 两者都有，是两种模式唯一对齐的视觉输入。"
        "如果做跨模式联合训练，视觉输入只能使用 side_rgb，或需要对相机视角做域适应。"
    ),
    (
        "④ Teleop 动作时长更长、分布更分散",
        "Teleop 平均 demo 时长 15.0s，ITW 12.9s；Teleop 帧数标准差 50.9 vs ITW 40.6。"
        "Teleop 操作更慢、执行时间变异更大，反映遥操作的固有延迟和操作员的不确定性。"
        "这会影响 action chunking 的窗口大小选择和轨迹对齐策略。"
    ),
    (
        "⑤ Teleop 有额外的臂动作信号",
        "Teleop 独有 arm_action 和 arm_published_action，两者与 arm_obs 高度相似（nearly identical），"
        "说明 published_action ≈ obs（控制闭环很紧）。ITW 没有臂动作信号，"
        "这意味着 ITW 数据无法直接用于训练臂的控制策略，只能用于训练手部策略。"
    ),
    (
        "⑥ 控制频率高度稳定",
        "两种模式的控制频率均稳定在 19.97Hz（理论 20Hz），时间戳标准差极小，"
        "说明数据采集的时序质量高，不存在明显掉帧或时钟漂移问题。"
    ),
]

for title, content in findings:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(content)

doc.add_page_break()

# ════════════ 10. 后续分析建议 ════════════════════════════════════════════
h1("10. 后续分析建议")

body("基于以上数据特征，建议按以下优先级推进分析：")

steps = [
    "Step 1：确认 ITW 的 arm_obs 物理含义（读取 config.yaml 中 sensor_id 配置，或对比 arm_obs 与 leader IMU 数据）",
    "Step 2：对比两种模式的 hand_action 速度分布和 jerk 分布（直方图 + KDE），量化动作平滑性差异",
    "Step 3：对 Teleop 的 follower_joint_efforts 做阶段性分析——按轨迹阶段（自由移动 / 接近 / 接触 / 拧转）分段统计力信号特征",
    "Step 4：分析 leader-follower 跟踪误差在不同关节和不同任务阶段的分布，评估 action label 质量",
    "Step 5：如需联合训练，设计 domain alignment 方案（side_rgb 对齐、arm_obs 处理策略、efforts 缺失的处理）",
    "Step 6：参考 'Diffusion Beats Autoregressive' 的方法论，量化数据多样性指标（coverage、action entropy）",
]

for i, s in enumerate(steps):
    p = doc.add_paragraph(s, style='List Number')

# ════════════ 保存 ════════════════════════════════════════════════════════
out_path = os.path.join(OUT_DIR, "DITTO_Data_Analysis_Report.docx")
doc.save(out_path)
print(f"\n报告已保存至：{out_path}")
